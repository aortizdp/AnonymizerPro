from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
import pandas as pd
import io
import zipfile
import traceback
from docx import Document
from backend.core.engine import AnonimitzadorEngine

app = FastAPI(title="AI Anonymizer API", version="13.0")
engine = AnonimitzadorEngine()

# --- FUNCIONS DE SUPORT ---
def load_file_agnostic(contents, filename):
    buffer = io.BytesIO(contents)
    fn = filename.lower()
    if fn.endswith(('.xlsx', '.xls')):
        return pd.read_excel(buffer, header=None), None
    
    # Per a CSV/TXT: Normalitzem línies i encodings
    for enc in ['mac_roman', 'utf-8', 'cp1252', 'latin-1']:
        try:
            text = contents.decode(enc).replace('\r\n', '\n').replace('\r', '\n')
            lines = [l for l in text.split('\n') if l.strip()]
            df = pd.DataFrame(lines)
            sep = ';' if text.count(';') > text.count(',') else ','
            return df, sep
        except: continue
    return pd.DataFrame([contents.decode('utf-8', errors='replace')]), ','

def save_file_agnostic(df, filename, sep=None):
    out = io.BytesIO()
    if filename.lower().endswith('.csv'):
        df.to_csv(out, index=False, header=False, sep=(sep if sep else ','), encoding='utf-8-sig')
        return out, "text/csv"
    df.to_excel(out, index=False, header=False)
    return out, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# --- ENDPOINT 1: ANONYMIZE ---
@app.post("/anonymize/")
async def anonymize_file(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        mapa_global = {}
        comptadors = {"USER": 0, "ID_CARD": 0, "IBAN": 0, "CREDIT_CARD": 0, "PHONE": 0, "EMAIL": 0}
        
        if file.filename.lower().endswith('.docx'):
            doc = Document(io.BytesIO(contents))
            for p in doc.paragraphs:
                if p.text.strip():
                    p.text, mapa_global, comptadors = engine.processar_text(p.text, mapa_global, comptadors)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text, mapa_global, comptadors = engine.processar_text(cell.text, mapa_global, comptadors)
            data_out = io.BytesIO()
            doc.save(data_out)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            df, sep = load_file_agnostic(contents, file.filename)
            for col in df.columns:
                df[col] = df[col].apply(lambda x: engine.processar_text(str(x), mapa_global, comptadors)[0] if pd.notna(x) else x)
            data_out, mime = save_file_agnostic(df, file.filename, sep)

        data_out.seek(0)
        df_map = pd.DataFrame(list(mapa_global.items()), columns=['Original', 'Token'])
        keys_out = io.BytesIO()
        df_map.to_excel(keys_out, index=False)
        keys_out.seek(0)

        zip_out = io.BytesIO()
        with zipfile.ZipFile(zip_out, "w") as zf:
            zf.writestr(f"ANONYMIZED_{file.filename}", data_out.getvalue())
            zf.writestr("decryption_keys.xlsx", keys_out.getvalue())
        zip_out.seek(0)

        return StreamingResponse(zip_out, media_type="application/zip", headers={"Content-Disposition": f"attachment; filename=ANONYMIZED_{file.filename}.zip"})
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=400, detail=str(e))

# --- ENDPOINT 2: DEANONYMIZE ---
@app.post("/deanonymize/")
async def deanonymize_file(
    file_anonim: UploadFile = File(...), 
    file_keys: UploadFile = File(...)
):
    try:
        k_cont = await file_keys.read()
        df_keys = pd.read_excel(io.BytesIO(k_cont))
        inv_map = dict(zip(df_keys['Token'], df_keys['Original']))

        a_cont = await file_anonim.read()
        out_buf = io.BytesIO()

        if file_anonim.filename.lower().endswith('.docx'):
            doc = Document(io.BytesIO(a_cont))
            for p in doc.paragraphs:
                p.text = engine.restaurar_text(p.text, inv_map)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = engine.restaurar_text(cell.text, inv_map)
            doc.save(out_buf)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            df, sep = load_file_agnostic(a_cont, file_anonim.filename)
            for col in df.columns:
                df[col] = df[col].apply(lambda x: engine.restaurar_text(str(x), inv_map) if pd.notna(x) else x)
            out_buf, mime = save_file_agnostic(df, file_anonim.filename, sep)

        out_buf.seek(0)
        return StreamingResponse(out_buf, media_type=mime, headers={"Content-Disposition": f"attachment; filename=RESTORED_{file_anonim.filename}"})
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=400, detail=str(e))
